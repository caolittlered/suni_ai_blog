"""
文档处理模块
"""

import os
from typing import List, Optional
from pathlib import Path
from langchain_core.documents import Document


def extract_text_from_file(file_path: str) -> str:
    """从文件中提取文本"""
    path = Path(file_path)
    suffix = path.suffix.lower()
    
    if suffix == '.txt' or suffix == '.md':
        with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
            return f.read()
    
    elif suffix == '.pdf':
        try:
            from pypdf import PdfReader
            reader = PdfReader(file_path)
            text = ""
            for page in reader.pages:
                text += page.extract_text() + "\n"
            return text
        except Exception as e:
            print(f"[DocProcessor] PDF 解析失败: {e}")
            return ""
    
    elif suffix == '.docx':
        try:
            from docx import Document as DocxDocument
            doc = DocxDocument(file_path)
            return "\n".join([para.text for para in doc.paragraphs])
        except Exception as e:
            print(f"[DocProcessor] DOCX 解析失败: {e}")
            return ""
    
    elif suffix == '.xlsx':
        try:
            from openpyxl import load_workbook
            wb = load_workbook(file_path)
            text = ""
            for sheet in wb.worksheets:
                for row in sheet.iter_rows(values_only=True):
                    text += " | ".join([str(cell) if cell else "" for cell in row]) + "\n"
            return text
        except Exception as e:
            print(f"[DocProcessor] XLSX 解析失败: {e}")
            return ""
    
    else:
        print(f"[DocProcessor] 不支持的文件类型: {suffix}")
        return ""


def process_file(file_path: str, user_id: int = None) -> Optional[Document]:
    """处理单个文件，返回 Document 对象"""
    path = Path(file_path)
    
    if not path.exists():
        return None
    
    text = extract_text_from_file(file_path)
    if not text.strip():
        return None
    
    metadata = {
        'source': str(path),
        'filename': path.name,
        'file_type': path.suffix.lower(),
    }
    if user_id:
        metadata['user_id'] = user_id
    
    return Document(page_content=text, metadata=metadata)


def process_directory(directory: str, user_id: int = None) -> List[Document]:
    """处理目录下的所有文档"""
    docs = []
    dir_path = Path(directory)
    
    if not dir_path.exists():
        print(f"[DocProcessor] 目录不存在: {directory}")
        return docs
    
    supported_extensions = {'.txt', '.md', '.pdf', '.docx', '.xlsx'}
    
    for file_path in dir_path.rglob('*'):
        if file_path.is_file() and file_path.suffix.lower() in supported_extensions:
            doc = process_file(str(file_path), user_id)
            if doc:
                docs.append(doc)
    
    print(f"[DocProcessor] 处理了 {len(docs)} 个文档")
    return docs